"""Shared pytest fixtures.

PDF and DocX fixtures are generated on the fly so no binary files need
to be committed to the repo.
"""

from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import pytest


@pytest.fixture()
def markdown_file(tmp_path: Path) -> Path:
    """A small Markdown document with headings, lists, and a code fence."""
    content = (
        "# Project Overview\n"
        "\n"
        "This document describes the parser.\n"
        "\n"
        "## Features\n"
        "\n"
        "- Fast parsing\n"
        "- Source tagging\n"
        "- Coordinates for PDF\n"
        "\n"
        "```python\n"
        "def hello():\n"
        "    return 'world'\n"
        "```\n"
        "\n"
        "Closing paragraph.\n"
    )
    path = tmp_path / "sample.md"
    path.write_text(content, encoding="utf-8")
    return path


@pytest.fixture()
def docx_file(tmp_path: Path) -> Path:
    """A DocX with a heading, body paragraphs, and a table."""
    document = DocxDocument()
    document.add_heading("DocX Report", level=1)

    paragraph = document.add_paragraph()
    run = paragraph.add_run("Intro paragraph with some body text.")
    run.font.size = Pt(11)

    document.add_heading("Data", level=2)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"

    centered = document.add_paragraph()
    centered.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    centered.add_run("A trailing paragraph.")

    path = tmp_path / "sample.docx"
    document.save(path)
    return path


@pytest.fixture()
def multi_page_pdf(tmp_path: Path) -> Path:
    """A born-digital 3-page PDF with headings and body text."""
    document = pymupdf.open()

    for page_index in range(3):
        page = document.new_page()
        title = f"Page {page_index + 1} Heading"
        body = (
            f"This is body text on page {page_index + 1}. "
            "It describes the content that follows. "
            "More sentences to form a decent sized paragraph."
        )
        page.insert_text((72, 72), title, fontsize=24)
        page.insert_text((72, 108), body, fontsize=12)

    path = tmp_path / "multi_page.pdf"
    document.save(path)
    document.close()
    return path


@pytest.fixture()
def empty_page_pdf(tmp_path: Path) -> Path:
    """A 2-page PDF where the second page is completely blank."""
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Only content lives here.", fontsize=12)
    document.new_page()  # blank page
    path = tmp_path / "empty_page.pdf"
    document.save(path)
    document.close()
    return path
