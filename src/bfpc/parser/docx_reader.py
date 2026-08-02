"""DocX reader backed by python-docx.

A DocX has no explicit paging model, so the document maps to a single
page whose blocks mirror the body element order (paragraphs and tables
interleaved as they appear in the file).
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from bfpc.parser.models import Block, BlockKind, Document, Page, Source


class DocxReader:
    """Parse a DocX file into a :class:`Document`."""

    def read(self, path: Path) -> Document:
        """Parse ``path`` and return its document model.

        :param path: path to a DocX file.
        :return: a single-page document.
        :raises FileNotFoundError: if the file does not exist.
        :raises ValueError: if the file is not a readable DocX.
        """
        path = Path(path)
        if not path.is_file():
            raise FileNotFoundError(f"DocX not found: {path}")

        try:
            document = docx.Document(path)
        except Exception as exc:
            raise ValueError(f"Not a valid DocX: {path}") from exc

        blocks = self._extract_blocks(document)
        return Document(
            path=path,
            source=Source.DOCX,
            pages=[Page(number=1, blocks=blocks)],
            metadata={"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
        )

    @staticmethod
    def _extract_blocks(document: DocxDocument) -> list[Block]:
        """Walk the body in document order, mixing paragraphs and tables."""
        blocks: list[Block] = []
        for element in document.element.body.iterchildren():
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, document)
                block = _paragraph_to_block(paragraph)
                if block is not None:
                    blocks.append(block)
            elif isinstance(element, CT_Tbl):
                table = Table(element, document)
                text = _table_to_text(table)
                if text:
                    blocks.append(Block(text=text, source=Source.DOCX, kind=BlockKind.TABLE))
        return blocks


def _paragraph_to_block(paragraph: Paragraph) -> Block | None:
    """Convert one paragraph to a Block, or None if it is empty."""
    text = paragraph.text.strip()
    if not text:
        return None

    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    kind = BlockKind.HEADING if style_name.lower().startswith("heading") else BlockKind.TEXT
    return Block(text=text, source=Source.DOCX, kind=kind)


def _table_to_text(table: Table) -> str:
    """Serialize a table into a readable multi-line string."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
